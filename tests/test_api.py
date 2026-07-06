import io
import json
import sqlite3
import tempfile
import unittest
from contextlib import closing
from pathlib import Path
from unittest.mock import Mock, patch

import requests
from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from langchain_core.documents import Document

from backend.app.config import Settings
from backend.app.main import create_app
from backend.app.models import (
    EnterpriseChatRequest,
    EnterpriseChatResponse,
    KnowledgeCitation,
    KnowledgeConfigUpdate,
    TursoConfigUpdate,
)
from backend.app.providers import BaseProvider, CompletionResult
from backend.app.providers.ollama_provider import OllamaProvider
from backend.app.services import (
    KnowledgeError,
    KnowledgeService,
    RuntimeConfigService,
    TursoError,
    TursoResult,
    TursoService,
)


ADMIN_EMAIL = "admin@example.com"
ADMIN_PASSWORD = "Senha-Admin-123!"


class FakeProvider(BaseProvider):
    name = "fake"

    def health_check(self) -> bool:
        return True

    def chat_completion(self, model, messages, temperature=0.2, max_tokens=None, structured=False):
        return CompletionResult(model=model, resposta="Resposta fundamentada [1].", metadados={"provider": "fake"})


class MemoryTursoService:
    """Substituto SQLite usado apenas para provar os contratos Turso nos testes."""

    configured = True

    def __init__(self, path: Path):
        self.path = path

    def initialize(self):
        with closing(self._connect()) as connection:
            connection.executescript(
                """
                CREATE TABLE IF NOT EXISTS users (
                    id TEXT PRIMARY KEY, email TEXT UNIQUE COLLATE NOCASE, display_name TEXT,
                    password_hash TEXT, role TEXT, is_active INTEGER, failed_attempts INTEGER DEFAULT 0,
                    locked_until TEXT, created_at TEXT, updated_at TEXT
                );
                CREATE TABLE IF NOT EXISTS sessions (
                    token_hash TEXT PRIMARY KEY, user_id TEXT, expires_at TEXT, created_at TEXT
                );
                CREATE TABLE IF NOT EXISTS conversations (
                    id TEXT PRIMARY KEY, user_id TEXT, title TEXT, created_at TEXT, updated_at TEXT
                );
                CREATE TABLE IF NOT EXISTS messages (
                    id TEXT PRIMARY KEY, conversation_id TEXT, role TEXT, content TEXT,
                    citations_json TEXT DEFAULT '[]', created_at TEXT
                );
                """
            )

    def execute(self, sql, args=()):
        try:
            with closing(self._connect()) as connection:
                cursor = connection.execute(sql, list(args))
                rows = [dict(row) for row in cursor.fetchall()] if cursor.description else []
                connection.commit()
                return TursoResult(rows=rows, affected_rows=max(cursor.rowcount, 0))
        except sqlite3.IntegrityError as exc:
            raise TursoError(409, "Registro duplicado no Turso.") from exc

    def transaction(self, statements):
        results = []
        with closing(self._connect()) as connection:
            for sql, args in statements:
                cursor = connection.execute(sql, args)
                rows = [dict(row) for row in cursor.fetchall()] if cursor.description else []
                results.append(TursoResult(rows=rows, affected_rows=max(cursor.rowcount, 0)))
            connection.commit()
        return results

    def test_connection(self):
        return {"status": "online", "database_url": "libsql://test.local", "encrypted": True}

    def _connect(self):
        connection = sqlite3.connect(self.path)
        connection.row_factory = sqlite3.Row
        return connection


class FakeKnowledge:
    def status(self):
        from backend.app.models import KnowledgeStatusResponse

        return KnowledgeStatusResponse(status="ready", enabled=True, configured=True, document_count=1)

    def ask(self, payload, conversation_id):
        citation = KnowledgeCitation(
            index=1,
            file_id="doc-1",
            file_name="Manual.md",
            file_type="MD",
            location="Documento completo",
            excerpt="Procedimento oficial.",
            web_url="https://drive.google.com/file/d/doc-1/view",
        )
        return EnterpriseChatResponse(
            status="ok",
            model="qwen3.5:9b",
            resposta="O procedimento oficial esta no manual [1].",
            conversation_id=conversation_id,
            citations=[citation],
            metadados={"source": "documents"},
        )


def make_settings(temp_dir: str, **overrides):
    base = Path(temp_dir)
    values = {
        "ENVIRONMENT": "development",
        "AUTH_ENABLED": False,
        "USER_AUTH_ENABLED": True,
        "REGISTRATION_ENABLED": True,
        "ADMIN_EMAIL": ADMIN_EMAIL,
        "LOCAL_CONFIG_DB_PATH": str(base / "config.db"),
        "CONFIG_KEY_PATH": str(base / ".config-key"),
        "KNOWLEDGE_INDEX_PATH": str(base / "knowledge"),
        "PROVIDER": "ollama",
        "DEFAULT_CHAT_MODEL": "qwen3.5:9b",
        "DEFAULT_VISION_MODEL": "qwen3.5:9b",
        "FRONTEND_ENABLED": False,
    }
    values.update(overrides)
    return Settings(**values)


class ApiSecurityAndHistoryTests(unittest.TestCase):
    def setUp(self):
        self.temp = tempfile.TemporaryDirectory()
        self.settings = make_settings(self.temp.name)
        self.storage = MemoryTursoService(Path(self.temp.name) / "turso-test.db")
        self.app = create_app(self.settings, turso_service=self.storage)
        self.app.state.provider = FakeProvider()
        self.app.state.knowledge_service = FakeKnowledge()
        self.app.state.assistant_service.knowledge = self.app.state.knowledge_service
        self.client = TestClient(self.app)

    def tearDown(self):
        self.client.close()
        self.temp.cleanup()

    @staticmethod
    def bearer(token):
        return {"Authorization": f"Bearer {token}"}

    def setup_admin(self):
        response = self.client.post(
            "/auth/setup",
            json={"email": ADMIN_EMAIL, "display_name": "Administrador", "password": ADMIN_PASSWORD},
        )
        self.assertEqual(response.status_code, 201, response.text)
        return response.json()

    def test_first_access_login_logout_and_password_hash_live_in_turso(self):
        status = self.client.get("/auth/status").json()
        self.assertTrue(status["storage_configured"])
        self.assertTrue(status["storage_online"])
        self.assertTrue(status["requires_setup"])

        created = self.setup_admin()
        self.assertEqual(created["user"]["role"], "admin")
        raw_turso = (Path(self.temp.name) / "turso-test.db").read_bytes()
        self.assertNotIn(ADMIN_PASSWORD.encode(), raw_turso)

        headers = self.bearer(created["access_token"])
        self.assertEqual(self.client.get("/auth/me", headers=headers).status_code, 200)
        self.assertEqual(self.client.post("/auth/logout", headers=headers).status_code, 204)
        self.assertEqual(self.client.get("/auth/me", headers=headers).status_code, 401)
        self.assertEqual(
            self.client.post("/auth/login", json={"email": ADMIN_EMAIL, "password": ADMIN_PASSWORD}).status_code,
            200,
        )

    def test_user_sees_only_chat_and_own_persisted_history(self):
        admin = self.setup_admin()
        registered = self.client.post(
            "/auth/register",
            json={"email": "analista@example.com", "display_name": "Analista", "password": "Senha-Analista-123!"},
        ).json()
        headers = self.bearer(registered["access_token"])
        self.assertEqual(self.client.get("/admin/config", headers=headers).status_code, 403)

        answer = self.client.post(
            "/assistant/chat",
            headers=headers,
            json={"pergunta": "Qual e o procedimento oficial?"},
        )
        self.assertEqual(answer.status_code, 200, answer.text)
        conversation_id = answer.json()["conversation_id"]
        conversations = self.client.get("/conversations", headers=headers).json()
        self.assertEqual(len(conversations), 1)
        detail = self.client.get(f"/conversations/{conversation_id}", headers=headers).json()
        self.assertEqual(len(detail["messages"]), 2)
        self.assertEqual(detail["messages"][1]["citations"][0]["file_name"], "Manual.md")

        admin_headers = self.bearer(admin["access_token"])
        self.assertEqual(self.client.get(f"/conversations/{conversation_id}", headers=admin_headers).status_code, 404)

    def test_turso_secret_is_admin_only_encrypted_and_never_returned(self):
        admin = self.setup_admin()
        token = "synthetic-turso-token-that-must-not-leak"
        response = self.client.put(
            "/admin/turso/config",
            headers=self.bearer(admin["access_token"]),
            json={"database_url": "libsql://aethra-test.turso.io", "auth_token": token},
        )
        self.assertEqual(response.status_code, 200, response.text)
        body = response.json()
        self.assertTrue(body["turso_token_configured"])
        self.assertNotIn("auth_token", body)
        raw_config = Path(self.settings.local_config_db_path).read_bytes()
        self.assertNotIn(token.encode(), raw_config)

    def test_legacy_analytical_routes_are_completely_removed(self):
        legacy_prefix = "/" + "d" + "w"
        paths = {route.path for route in self.app.routes}
        self.assertFalse(any(path.startswith(legacy_prefix) for path in paths))
        self.assertEqual(self.client.get(f"{legacy_prefix}/schema").status_code, 404)

    def test_cors_allows_authorization_and_delete(self):
        response = self.client.options(
            "/conversations/test",
            headers={
                "Origin": "http://localhost:5500",
                "Access-Control-Request-Method": "DELETE",
                "Access-Control-Request-Headers": "Authorization",
            },
        )
        self.assertEqual(response.status_code, 200)
        self.assertIn("DELETE", response.headers["access-control-allow-methods"])


class RuntimeAndTursoClientTests(unittest.TestCase):
    def setUp(self):
        self.temp = tempfile.TemporaryDirectory()
        self.settings = make_settings(self.temp.name)
        self.runtime = RuntimeConfigService(self.settings)
        self.runtime.initialize()

    def tearDown(self):
        self.temp.cleanup()

    @patch("backend.app.services.turso_service.requests.post")
    def test_turso_client_uses_bound_parameters_and_bearer_auth(self, post):
        self.runtime.save_turso_config(
            TursoConfigUpdate(database_url="libsql://test-db.turso.io", auth_token="synthetic-token")
        )
        response = Mock()
        response.raise_for_status.return_value = None
        response.json.return_value = {
            "results": [
                {
                    "type": "ok",
                    "response": {
                        "result": {
                            "cols": [{"name": "email"}],
                            "rows": [[{"type": "text", "value": "user@example.com"}]],
                            "affected_row_count": 0,
                        }
                    },
                },
                {"type": "ok", "response": {"type": "close"}},
            ]
        }
        post.return_value = response
        service = TursoService(self.runtime)

        result = service.execute("SELECT email FROM users WHERE id = ?", ["user-1"])

        self.assertEqual(result.rows[0]["email"], "user@example.com")
        request_json = post.call_args.kwargs["json"]
        self.assertEqual(request_json["requests"][0]["stmt"]["args"][0]["value"], "user-1")
        self.assertEqual(post.call_args.kwargs["headers"]["Authorization"], "Bearer synthetic-token")

    def test_drive_and_turso_credentials_are_both_encrypted(self):
        turso_token = "secret-turso"
        drive_key = "secret-drive-private-key"
        self.runtime.save_turso_config(
            TursoConfigUpdate(database_url="libsql://test-db.turso.io", auth_token=turso_token)
        )
        self.runtime.save_knowledge_config(
            KnowledgeConfigUpdate(
                folder_id="folder-test-123",
                service_account_json=json.dumps({
                    "type": "service_account",
                    "client_email": "reader@example.iam.gserviceaccount.com",
                    "private_key": drive_key,
                    "token_uri": "https://oauth2.googleapis.com/token",
                }),
            ),
            updated_by="admin",
        )
        raw = Path(self.settings.local_config_db_path).read_bytes()
        self.assertNotIn(turso_token.encode(), raw)
        self.assertNotIn(drive_key.encode(), raw)


class KnowledgeServiceTests(unittest.TestCase):
    def setUp(self):
        self.temp = tempfile.TemporaryDirectory()
        self.settings = make_settings(self.temp.name)
        self.runtime = RuntimeConfigService(self.settings)
        self.runtime.initialize()
        self.runtime.save_knowledge_config(
            KnowledgeConfigUpdate(
                folder_id="folder-test-123",
                service_account_json=json.dumps({
                    "type": "service_account",
                    "client_email": "reader@example.iam.gserviceaccount.com",
                    "private_key": "synthetic-key",
                    "token_uri": "https://oauth2.googleapis.com/token",
                }),
            ),
            updated_by="admin",
        )
        self.runtime.save_knowledge_status(
            {"status": "ready", "document_count": 2, "page_count": 2, "chunk_count": 4},
            "admin",
        )
        self.service = KnowledgeService(FakeProvider(), self.settings, self.runtime)

    def tearDown(self):
        self.temp.cleanup()

    def test_extracts_markdown_text_and_docx(self):
        markdown = self.service._extract_documents(
            b"# Politica\nPrazo de 30 dias.",
            {"id": "md-1", "name": "politica.md", "mimeType": "text/markdown"},
        )
        docx = DocxDocument()
        docx.add_paragraph("Manual interno de atendimento.")
        stream = io.BytesIO()
        docx.save(stream)
        word = self.service._extract_documents(
            stream.getvalue(),
            {"id": "docx-1", "name": "manual.docx", "mimeType": self.service.DOCX_MIME},
        )
        self.assertIn("Prazo de 30 dias", markdown[0].page_content)
        self.assertEqual(markdown[0].metadata["file_type"], "MD")
        self.assertIn("Manual interno", word[0].page_content)
        self.assertEqual(word[0].metadata["file_type"], "DOCX")

    def test_answer_contains_document_location_and_citation(self):
        store = Mock()
        source = Document(
            page_content="O prazo contratual e de 30 dias.",
            metadata={
                "file_id": "pdf-123",
                "file_name": "Contrato.pdf",
                "file_type": "PDF",
                "page": 7,
                "location": "Pagina 7",
                "web_url": "https://drive.google.com/file/d/pdf-123/view",
            },
        )
        store.similarity_search.return_value = [source]
        store.max_marginal_relevance_search.return_value = [source]
        self.service._vector_store = lambda config: store
        result = self.service.ask(EnterpriseChatRequest(pergunta="Qual e o prazo?"), "conversation-1")
        self.assertEqual(result.citations[0].page, 7)
        self.assertEqual(result.citations[0].location, "Pagina 7")
        self.assertIn("[1]", result.resposta)
        self.assertEqual(result.metadados["retrieval_strategy"], "relevance_plus_mmr")

    def test_broad_retrieval_prioritizes_multiple_files(self):
        def source(file_id, page, text):
            return Document(
                page_content=text,
                metadata={"file_id": file_id, "location": f"Pagina {page}"},
            )

        file_a = [source("a", page, f"Atendimento processo cliente etapa {page}") for page in range(1, 6)]
        file_b = source("b", 1, "Atendimento indicadores e qualidade")
        file_c = source("c", 1, "Atendimento riscos e responsabilidades")
        selected = self.service._fuse_retrieval(
            "Correlacione os principais temas dos documentos sobre atendimento",
            [*file_a, file_b, file_c],
            [file_b, file_c, file_a[0]],
        )
        file_ids = [item.metadata["file_id"] for item in selected]
        self.assertEqual(set(file_ids), {"a", "b", "c"})
        self.assertLessEqual(file_ids.count("a"), 2)

    @patch(
        "backend.app.services.knowledge_service.requests.post",
        side_effect=requests.ConnectionError("offline"),
    )
    def test_sync_reports_when_ollama_is_offline(self, _post):
        with self.assertRaises(KnowledgeError) as context:
            self.service._ensure_embedding_runtime(self.runtime.get_knowledge_config())
        self.assertEqual(context.exception.status_code, 503)
        self.assertIn("Ollama esta offline", context.exception.detail)

    def test_rebuild_index_sends_small_embedding_batches(self):
        store = Mock()
        self.service._vector_store = lambda _config: store
        chunks = [
            Document(
                page_content=f"Trecho documental {index}",
                metadata={"file_id": "pdf-1", "location": f"Pagina {index + 1}"},
            )
            for index in range(35)
        ]
        self.service._rebuild_index(chunks, self.runtime.get_knowledge_config())
        batch_sizes = [len(call.args[0]) for call in store.add_documents.call_args_list]
        self.assertEqual(batch_sizes, [16, 16, 3])


class OllamaReliabilityTests(unittest.TestCase):
    @patch("backend.app.providers.ollama_provider.requests.post")
    def test_qwen_disables_thinking_to_preserve_final_answer(self, post):
        response = Mock()
        response.raise_for_status.return_value = None
        response.json.return_value = {"response": "OK", "done": True}
        post.return_value = response
        provider = OllamaProvider("http://127.0.0.1:11434", timeout=30)
        result = provider.chat_completion("qwen3.5:9b", [{"role": "user", "content": "Responda OK"}], structured=True)
        self.assertFalse(post.call_args.kwargs["json"]["think"])
        self.assertEqual(result.resposta, "OK")


if __name__ == "__main__":
    unittest.main()
